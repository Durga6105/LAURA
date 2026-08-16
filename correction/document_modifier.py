"""
LAURA - Document Modifier

Applies approved clause modifications to NDA documents
without overwriting the original file.

Supported:
    DOCX
    PDF
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import pymupdf
from docx import Document


class DocumentModifier:
    """Apply clause-level modifications to NDA documents."""

    SUPPORTED_EXTENSIONS = {
        ".docx",
        ".pdf",
    }

    def modify(
        self,
        input_path: str | Path,
        modifications: list[dict[str, Any]],
        output_path: str | Path,
    ) -> Path:
        """
        Create a modified NDA.

        The original document is never overwritten.
        """

        source = Path(input_path)
        destination = Path(output_path)

        if not source.exists():
            raise FileNotFoundError(
                f"NDA not found: {source}"
            )

        if source.resolve() == destination.resolve():
            raise ValueError(
                "Output document must not overwrite the original."
            )

        extension = source.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported document type: {extension}"
            )

        destination.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        if extension == ".docx":
            return self._modify_docx(
                source,
                destination,
                modifications,
            )

        return self._modify_pdf(
            source,
            destination,
            modifications,
        )

    # =========================================================
    # DOCX
    # =========================================================

    def _modify_docx(
        self,
        source: Path,
        destination: Path,
        modifications: list[dict[str, Any]],
    ) -> Path:
        """Apply text replacements to a DOCX document."""

        document = Document(source)

        replacement_count = 0

        for modification in modifications:

            original_text = str(
                modification.get(
                    "original_text",
                    "",
                )
            ).strip()

            modified_text = str(
                modification.get(
                    "modified_text",
                    "",
                )
            ).strip()

            if not original_text or not modified_text:
                continue

            replaced = False

            # -------------------------------------------------
            # Search normal paragraphs
            # -------------------------------------------------

            for paragraph in document.paragraphs:

                if original_text in paragraph.text:

                    self._replace_in_paragraph(
                        paragraph,
                        original_text,
                        modified_text,
                    )

                    replacement_count += 1
                    replaced = True
                    break

            if replaced:
                continue

            # -------------------------------------------------
            # Search tables
            # -------------------------------------------------

            for table in document.tables:

                if self._replace_in_table(
                    table,
                    original_text,
                    modified_text,
                ):
                    replacement_count += 1
                    replaced = True
                    break

        if replacement_count == 0:
            raise ValueError(
                "No matching NDA clauses were found "
                "for the requested modifications."
            )

        document.save(destination)

        return destination

    @staticmethod
    def _replace_in_paragraph(
        paragraph: Any,
        original_text: str,
        modified_text: str,
    ) -> None:
        """Replace text within a DOCX paragraph."""

        full_text = paragraph.text

        new_text = full_text.replace(
            original_text,
            modified_text,
            1,
        )

        # Clear existing runs.
        for run in paragraph.runs:
            run.text = ""

        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)

    def _replace_in_table(
        self,
        table: Any,
        original_text: str,
        modified_text: str,
    ) -> bool:
        """Replace text inside DOCX table cells."""

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    if original_text in paragraph.text:

                        self._replace_in_paragraph(
                            paragraph,
                            original_text,
                            modified_text,
                        )

                        return True

        return False

    # =========================================================
    # PDF
    # =========================================================

    def _modify_pdf(
        self,
        source: Path,
        destination: Path,
        modifications: list[dict[str, Any]],
    ) -> Path:
        """
        Apply clause-level replacements to a PDF.

        PyMuPDF is used to:
            1. Locate the original clause.
            2. Redact the original text.
            3. Insert the corrected clause.
            4. Save a new PDF.

        The original PDF is never modified.
        """

        document = pymupdf.open(source)

        replacement_count = 0

        try:

            for modification in modifications:

                original_text = str(
                    modification.get(
                        "original_text",
                        "",
                    )
                ).strip()

                modified_text = str(
                    modification.get(
                        "modified_text",
                        "",
                    )
                ).strip()

                if not original_text or not modified_text:
                    continue

                page_number = modification.get(
                    "page_number"
                )

                # -------------------------------------------------
                # If validation knows the page, search there first.
                # Page numbers are 1-based in LAURA.
                # PyMuPDF pages are 0-based.
                # -------------------------------------------------

                pages_to_search = []

                if page_number:

                    try:
                        page_index = int(
                            page_number
                        ) - 1

                        if (
                            0
                            <= page_index
                            < len(document)
                        ):
                            pages_to_search = [
                                page_index
                            ]

                    except (
                        TypeError,
                        ValueError,
                    ):
                        pages_to_search = []

                # If page number isn't available, search entire PDF.
                if not pages_to_search:
                    pages_to_search = list(
                        range(len(document))
                    )

                found = False

                for page_index in pages_to_search:

                    page = document[
                        page_index
                    ]

                    matches = page.search_for(
                        original_text
                    )

                    if not matches:
                        continue

                    # -------------------------------------------------
                    # A clause can span multiple text rectangles.
                    # Combine all matching rectangles into one region.
                    # -------------------------------------------------

                    replacement_rect = (
                        self._combine_rects(
                            matches
                        )
                    )

                    # -------------------------------------------------
                    # Redact original text.
                    # -------------------------------------------------

                    page.add_redact_annot(
                        replacement_rect,
                        fill=(1, 1, 1),
                    )

                    page.apply_redactions()

                    # -------------------------------------------------
                    # Determine a reasonable font size.
                    # -------------------------------------------------

                    font_size = self._calculate_font_size(
                        replacement_rect,
                        modified_text,
                    )

                    # -------------------------------------------------
                    # Insert corrected text.
                    # -------------------------------------------------

                    inserted = page.insert_textbox(
                        replacement_rect,
                        modified_text,
                        fontsize=font_size,
                        fontname="helv",
                        color=(0, 0, 0),
                        align=0,
                    )

                    # PyMuPDF returns a negative value when
                    # the text does not fit inside the rectangle.
                    if inserted < 0:

                        # Try progressively smaller fonts.
                        inserted = self._insert_with_fallback(
                            page=page,
                            rect=replacement_rect,
                            text=modified_text,
                            initial_font_size=font_size,
                        )

                    if inserted < 0:
                        raise ValueError(
                            "Corrected clause is too large "
                            "to fit in the original PDF text area."
                        )

                    replacement_count += 1
                    found = True

                    break

                if not found:
                    raise ValueError(
                        "Could not find the original NDA "
                        "clause in the PDF:\n"
                        f"{original_text}"
                    )

            if replacement_count == 0:
                raise ValueError(
                    "No matching NDA clauses were found "
                    "for the requested PDF modifications."
                )

            document.save(
                destination,
                garbage=4,
                deflate=True,
            )

        finally:
            document.close()

        return destination

    # =========================================================
    # PDF Helpers
    # =========================================================

    @staticmethod
    def _combine_rects(
        rectangles: list[pymupdf.Rect],
    ) -> pymupdf.Rect:
        """Combine multiple text rectangles."""

        if not rectangles:
            raise ValueError(
                "No PDF text rectangles found."
            )

        combined = pymupdf.Rect(
            rectangles[0]
        )

        for rect in rectangles[1:]:
            combined |= pymupdf.Rect(rect)

        return combined

    @staticmethod
    def _calculate_font_size(
        rect: pymupdf.Rect,
        text: str,
    ) -> float:
        """
        Estimate a suitable PDF font size.

        The replacement should initially use the
        approximate height of the original text area.
        """

        height = rect.height

        if height <= 0:
            return 10.0

        # Typical body text is roughly 65-75% of
        # the text rectangle height.
        font_size = height * 0.70

        # Keep the initial size within sensible bounds.
        font_size = max(
            7.0,
            min(font_size, 14.0),
        )

        return font_size

    @staticmethod
    def _insert_with_fallback(
        page: pymupdf.Page,
        rect: pymupdf.Rect,
        text: str,
        initial_font_size: float,
    ) -> float:
        """
        Retry PDF text insertion using smaller fonts.
        """

        font_sizes = [
            initial_font_size * 0.90,
            initial_font_size * 0.80,
            initial_font_size * 0.70,
            initial_font_size * 0.60,
            7.0,
        ]

        for font_size in font_sizes:

            result = page.insert_textbox(
                rect,
                text,
                fontsize=max(
                    font_size,
                    6.0,
                ),
                fontname="helv",
                color=(0, 0, 0),
                align=0,
            )

            if result >= 0:
                return result

        return -1.0