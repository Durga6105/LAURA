"""
LAURA - Report Generator

Generates NDA validation reports containing:

- Overall PASS/FAIL status
- Overall risk
- Rule-by-rule validation results
- Failed rules
- Corrections made
- Re-validation history

Supported formats:
- JSON
- DOCX
- PDF
"""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    SimpleDocTemplate,
    KeepTogether,
)


class ReportGenerator:
    """Generate NDA validation reports."""

    # ======================================================
    # REPORT DATA
    # ======================================================

    def generate_report_data(
        self,
        document_name: str,
        final_status: str,
        overall_risk: str,
        validation_results: list[dict[str, Any]],
        modifications: list[dict[str, Any]] | None = None,
        validation_history: list[dict[str, Any]] | None = None,
    ) -> dict[str, Any]:
        """
        Build the complete report data structure.
        """

        modifications = modifications or []
        validation_history = validation_history or []

        passed = sum(
            result.get("status") == "PASS"
            for result in validation_results
        )

        failed = sum(
            result.get("status") == "FAIL"
            for result in validation_results
        )

        not_found = sum(
            result.get("status") == "NOT_FOUND"
            for result in validation_results
        )

        uncertain = sum(
            result.get("status") == "UNCERTAIN"
            for result in validation_results
        )

        high_risk = sum(
            result.get("risk") == "HIGH"
            and result.get("status") != "PASS"
            for result in validation_results
        )

        medium_risk = sum(
            result.get("risk") == "MEDIUM"
            and result.get("status") != "PASS"
            for result in validation_results
        )

        low_risk = sum(
            result.get("risk") == "LOW"
            and result.get("status") != "PASS"
            for result in validation_results
        )

        return {
            "report_metadata": {
                "application": "LAURA",
                "report_type": "NDA Validation Report",
                "generated_at": datetime.now().isoformat(),
                "document_name": document_name,
            },
            "overall_result": {
                "status": final_status,
                "risk": overall_risk,
            },
            "statistics": {
                "total_rules": len(validation_results),
                "passed": passed,
                "failed": failed,
                "not_found": not_found,
                "uncertain": uncertain,
                "high_risk_issues": high_risk,
                "medium_risk_issues": medium_risk,
                "low_risk_issues": low_risk,
                "modifications": len(modifications),
                "validation_iterations": len(
                    validation_history
                ),
            },
            "rule_results": validation_results,
            "modifications": modifications,
            "validation_history": validation_history,
        }

    # ======================================================
    # JSON
    # ======================================================

    def save_json(
        self,
        report_data: dict[str, Any],
        output_path: str | Path,
    ) -> Path:
        """Save report as JSON."""

        output = Path(output_path)

        output.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        output.write_text(
            json.dumps(
                report_data,
                indent=4,
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        return output

    # ======================================================
    # DOCX
    # ======================================================

    def save_docx(
        self,
        report_data: dict[str, Any],
        output_path: str | Path,
    ) -> Path:
        """Save a human-readable report as DOCX."""

        output = Path(output_path)

        output.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        document = Document()

        # --------------------------------------------------
        # Title
        # --------------------------------------------------

        title = document.add_heading(
            "LAURA - NDA Validation Report",
            level=0,
        )

        title.runs[0].font.size = Pt(22)

        metadata = report_data.get(
            "report_metadata",
            {},
        )

        document.add_paragraph(
            f"Document: "
            f"{metadata.get('document_name', '')}"
        )

        document.add_paragraph(
            f"Generated: "
            f"{metadata.get('generated_at', '')}"
        )

        # --------------------------------------------------
        # Overall Result
        # --------------------------------------------------

        document.add_heading(
            "1. Overall Result",
            level=1,
        )

        overall = report_data.get(
            "overall_result",
            {},
        )

        document.add_paragraph(
            f"Status: {overall.get('status', '')}"
        )

        document.add_paragraph(
            f"Overall Risk: {overall.get('risk', '')}"
        )

        # --------------------------------------------------
        # Statistics
        # --------------------------------------------------

        document.add_heading(
            "2. Validation Summary",
            level=1,
        )

        statistics = report_data.get(
            "statistics",
            {},
        )

        summary_table = document.add_table(
            rows=1,
            cols=2,
        )

        summary_table.style = "Table Grid"

        summary_table.rows[0].cells[0].text = "Metric"
        summary_table.rows[0].cells[1].text = "Value"

        summary_items = [
            ("Total Rules", statistics.get("total_rules", 0)),
            ("Passed", statistics.get("passed", 0)),
            ("Failed", statistics.get("failed", 0)),
            ("Not Found", statistics.get("not_found", 0)),
            ("Uncertain", statistics.get("uncertain", 0)),
            (
                "High Risk Issues",
                statistics.get("high_risk_issues", 0),
            ),
            (
                "Medium Risk Issues",
                statistics.get("medium_risk_issues", 0),
            ),
            (
                "Low Risk Issues",
                statistics.get("low_risk_issues", 0),
            ),
            (
                "Modifications",
                statistics.get("modifications", 0),
            ),
            (
                "Validation Iterations",
                statistics.get(
                    "validation_iterations",
                    0,
                ),
            ),
        ]

        for metric, value in summary_items:
            row = summary_table.add_row()
            row.cells[0].text = str(metric)
            row.cells[1].text = str(value)

        # --------------------------------------------------
        # Rule Results
        # --------------------------------------------------

        document.add_heading(
            "3. Rule-by-Rule Results",
            level=1,
        )

        results = report_data.get(
            "rule_results",
            [],
        )

        for result in results:

            rule_id = result.get(
                "rule_id",
                "",
            )

            status = result.get(
                "status",
                "",
            )

            risk = result.get(
                "risk",
                "",
            )

            document.add_heading(
                f"Rule {rule_id}",
                level=2,
            )

            document.add_paragraph(
                f"Status: {status}"
            )

            document.add_paragraph(
                f"Risk: {risk}"
            )

            document.add_paragraph(
                f"Mandatory: "
                f"{result.get('mandatory', '')}"
            )

            document.add_paragraph(
                f"Evidence: "
                f"{result.get('evidence', '')}"
            )

            document.add_paragraph(
                f"Reason: "
                f"{result.get('reason', '')}"
            )

            required_change = result.get(
                "required_change",
                "",
            )

            if required_change:
                document.add_paragraph(
                    f"Required Change: "
                    f"{required_change}"
                )

        # --------------------------------------------------
        # Modifications
        # --------------------------------------------------

        document.add_heading(
            "4. Modifications Made",
            level=1,
        )

        modifications = report_data.get(
            "modifications",
            [],
        )

        if not modifications:

            document.add_paragraph(
                "No modifications were made."
            )

        else:

            for index, modification in enumerate(
                modifications,
                start=1,
            ):

                document.add_heading(
                    f"Modification {index}",
                    level=2,
                )

                document.add_paragraph(
                    f"Rule ID: "
                    f"{modification.get('rule_id', '')}"
                )

                document.add_paragraph(
                    "Original Text:"
                )

                document.add_paragraph(
                    modification.get(
                        "original_text",
                        "",
                    )
                )

                document.add_paragraph(
                    "Modified Text:"
                )

                document.add_paragraph(
                    modification.get(
                        "modified_text",
                        "",
                    )
                )

                document.add_paragraph(
                    f"Reason: "
                    f"{modification.get('reason', '')}"
                )

        # --------------------------------------------------
        # Validation History
        # --------------------------------------------------

        document.add_heading(
            "5. Validation History",
            level=1,
        )

        history = report_data.get(
            "validation_history",
            [],
        )

        if not history:

            document.add_paragraph(
                "No validation history available."
            )

        else:

            for iteration in history:

                document.add_paragraph(
                    f"Iteration: "
                    f"{iteration.get('iteration', '')}"
                )

                document.add_paragraph(
                    f"Status: "
                    f"{iteration.get('status', '')}"
                )

                document.add_paragraph(
                    f"Risk: "
                    f"{iteration.get('risk', '')}"
                )

                document.add_paragraph(
                    f"Passed Rules: "
                    f"{iteration.get('passed_rules', '')}"
                )

                document.add_paragraph(
                    f"Failed Rules: "
                    f"{iteration.get('failed_rules', '')}"
                )

        document.save(output)

        return output

    # ======================================================
    # PDF
    # ======================================================

    def save_pdf(
        self,
        report_data: dict[str, Any],
        output_path: str | Path,
    ) -> Path:
        """
        Save a human-readable NDA validation report as PDF.
        """

        output = Path(output_path)

        output.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "LAURAReportTitle",
            parent=styles["Title"],
            fontSize=20,
            leading=24,
            alignment=TA_CENTER,
            spaceAfter=12,
        )

        heading_style = ParagraphStyle(
            "LAURAHeading",
            parent=styles["Heading2"],
            fontSize=13,
            leading=16,
            spaceBefore=10,
            spaceAfter=6,
        )

        subheading_style = ParagraphStyle(
            "LAURASubHeading",
            parent=styles["Heading3"],
            fontSize=11,
            leading=14,
            spaceBefore=7,
            spaceAfter=4,
        )

        body_style = ParagraphStyle(
            "LAURABody",
            parent=styles["BodyText"],
            fontSize=9,
            leading=13,
            spaceAfter=5,
        )

        small_style = ParagraphStyle(
            "LAURASmall",
            parent=styles["BodyText"],
            fontSize=8,
            leading=11,
            spaceAfter=3,
        )

        document = SimpleDocTemplate(
            str(output),
            pagesize=A4,
            rightMargin=15 * mm,
            leftMargin=15 * mm,
            topMargin=15 * mm,
            bottomMargin=15 * mm,
            title="LAURA - NDA Validation Report",
            author="LAURA",
        )

        story: list[Any] = []

        metadata = report_data.get(
            "report_metadata",
            {},
        )

        overall = report_data.get(
            "overall_result",
            {},
        )

        statistics = report_data.get(
            "statistics",
            {},
        )

        # --------------------------------------------------
        # Title
        # --------------------------------------------------

        story.append(
            Paragraph(
                "LAURA - NDA Validation Report",
                title_style,
            )
        )

        story.append(
            Paragraph(
                f"<b>Document:</b> "
                f"{self._escape(metadata.get('document_name', ''))}",
                body_style,
            )
        )

        story.append(
            Paragraph(
                f"<b>Generated:</b> "
                f"{self._escape(metadata.get('generated_at', ''))}",
                small_style,
            )
        )

        story.append(
            Spacer(
                1,
                8,
            )
        )

        # --------------------------------------------------
        # Overall Result
        # --------------------------------------------------

        story.append(
            Paragraph(
                "1. Overall Result",
                heading_style,
            )
        )

        status = self._escape(
            overall.get("status", "")
        )

        risk = self._escape(
            overall.get("risk", "")
        )

        overall_table = Table(
            [
                ["Status", status],
                ["Overall Risk", risk],
            ],
            colWidths=[
                45 * mm,
                120 * mm,
            ],
        )

        overall_table.setStyle(
            TableStyle(
                [
                    (
                        "GRID",
                        (0, 0),
                        (-1, -1),
                        0.5,
                        colors.grey,
                    ),
                    (
                        "FONTNAME",
                        (0, 0),
                        (-1, -1),
                        "Helvetica",
                    ),
                    (
                        "FONTNAME",
                        (0, 0),
                        (0, -1),
                        "Helvetica-Bold",
                    ),
                    (
                        "VALIGN",
                        (0, 0),
                        (-1, -1),
                        "TOP",
                    ),
                    (
                        "LEFTPADDING",
                        (0, 0),
                        (-1, -1),
                        6,
                    ),
                    (
                        "RIGHTPADDING",
                        (0, 0),
                        (-1, -1),
                        6,
                    ),
                    (
                        "TOPPADDING",
                        (0, 0),
                        (-1, -1),
                        6,
                    ),
                    (
                        "BOTTOMPADDING",
                        (0, 0),
                        (-1, -1),
                        6,
                    ),
                ]
            )
        )

        story.append(
            overall_table
        )

        # --------------------------------------------------
        # Validation Summary
        # --------------------------------------------------

        story.append(
            Paragraph(
                "2. Validation Summary",
                heading_style,
            )
        )

        summary_items = [
            (
                "Total Rules",
                statistics.get(
                    "total_rules",
                    0,
                ),
            ),
            (
                "Passed",
                statistics.get(
                    "passed",
                    0,
                ),
            ),
            (
                "Failed",
                statistics.get(
                    "failed",
                    0,
                ),
            ),
            (
                "Not Found",
                statistics.get(
                    "not_found",
                    0,
                ),
            ),
            (
                "Uncertain",
                statistics.get(
                    "uncertain",
                    0,
                ),
            ),
            (
                "High Risk Issues",
                statistics.get(
                    "high_risk_issues",
                    0,
                ),
            ),
            (
                "Medium Risk Issues",
                statistics.get(
                    "medium_risk_issues",
                    0,
                ),
            ),
            (
                "Low Risk Issues",
                statistics.get(
                    "low_risk_issues",
                    0,
                ),
            ),
            (
                "Modifications",
                statistics.get(
                    "modifications",
                    0,
                ),
            ),
            (
                "Validation Iterations",
                statistics.get(
                    "validation_iterations",
                    0,
                ),
            ),
        ]

        summary_data = [
            [
                Paragraph(
                    "<b>Metric</b>",
                    small_style,
                ),
                Paragraph(
                    "<b>Value</b>",
                    small_style,
                ),
            ]
        ]

        for metric, value in summary_items:
            summary_data.append(
                [
                    Paragraph(
                        self._escape(
                            str(metric)
                        ),
                        small_style,
                    ),
                    Paragraph(
                        self._escape(
                            str(value)
                        ),
                        small_style,
                    ),
                ]
            )

        summary_table = Table(
            summary_data,
            colWidths=[
                120 * mm,
                45 * mm,
            ],
            repeatRows=1,
        )

        summary_table.setStyle(
            TableStyle(
                [
                    (
                        "GRID",
                        (0, 0),
                        (-1, -1),
                        0.5,
                        colors.grey,
                    ),
                    (
                        "BACKGROUND",
                        (0, 0),
                        (-1, 0),
                        colors.lightgrey,
                    ),
                    (
                        "VALIGN",
                        (0, 0),
                        (-1, -1),
                        "TOP",
                    ),
                    (
                        "LEFTPADDING",
                        (0, 0),
                        (-1, -1),
                        6,
                    ),
                    (
                        "RIGHTPADDING",
                        (0, 0),
                        (-1, -1),
                        6,
                    ),
                    (
                        "TOPPADDING",
                        (0, 0),
                        (-1, -1),
                        4,
                    ),
                    (
                        "BOTTOMPADDING",
                        (0, 0),
                        (-1, -1),
                        4,
                    ),
                ]
            )
        )

        story.append(
            summary_table
        )

        # --------------------------------------------------
        # Rule Results
        # --------------------------------------------------

        story.append(
            Paragraph(
                "3. Rule-by-Rule Results",
                heading_style,
            )
        )

        results = report_data.get(
            "rule_results",
            [],
        )

        if not results:

            story.append(
                Paragraph(
                    "No rule validation results available.",
                    body_style,
                )
            )

        else:

            for result in results:

                rule_id = self._escape(
                    result.get(
                        "rule_id",
                        "",
                    )
                )

                status = self._escape(
                    result.get(
                        "status",
                        "",
                    )
                )

                risk = self._escape(
                    result.get(
                        "risk",
                        "",
                    )
                )

                mandatory = self._escape(
                    str(
                        result.get(
                            "mandatory",
                            "",
                        )
                    )
                )

                evidence = self._escape(
                    result.get(
                        "evidence",
                        "",
                    )
                )

                reason = self._escape(
                    result.get(
                        "reason",
                        "",
                    )
                )

                required_change = self._escape(
                    result.get(
                        "required_change",
                        "",
                    )
                )

                rule_block: list[Any] = [
                    Paragraph(
                        f"Rule {rule_id}",
                        subheading_style,
                    ),
                    Paragraph(
                        f"<b>Status:</b> {status}",
                        body_style,
                    ),
                    Paragraph(
                        f"<b>Risk:</b> {risk}",
                        body_style,
                    ),
                    Paragraph(
                        f"<b>Mandatory:</b> {mandatory}",
                        body_style,
                    ),
                    Paragraph(
                        f"<b>Evidence:</b> {evidence}",
                        body_style,
                    ),
                    Paragraph(
                        f"<b>Reason:</b> {reason}",
                        body_style,
                    ),
                ]

                if required_change:
                    rule_block.append(
                        Paragraph(
                            f"<b>Required Change:</b> "
                            f"{required_change}",
                            body_style,
                        )
                    )

                story.append(
                    KeepTogether(
                        rule_block
                    )
                )

        # --------------------------------------------------
        # Modifications
        # --------------------------------------------------

        story.append(
            Paragraph(
                "4. Modifications Made",
                heading_style,
            )
        )

        modifications = report_data.get(
            "modifications",
            [],
        )

        if not modifications:

            story.append(
                Paragraph(
                    "No modifications were made.",
                    body_style,
                )
            )

        else:

            for index, modification in enumerate(
                modifications,
                start=1,
            ):

                rule_id = self._escape(
                    modification.get(
                        "rule_id",
                        "",
                    )
                )

                original_text = self._escape(
                    modification.get(
                        "original_text",
                        "",
                    )
                )

                modified_text = self._escape(
                    modification.get(
                        "modified_text",
                        "",
                    )
                )

                reason = self._escape(
                    modification.get(
                        "reason",
                        "",
                    )
                )

                modification_block = [
                    Paragraph(
                        f"Modification {index}",
                        subheading_style,
                    ),
                    Paragraph(
                        f"<b>Rule ID:</b> {rule_id}",
                        body_style,
                    ),
                    Paragraph(
                        "<b>Original Text:</b>",
                        body_style,
                    ),
                    Paragraph(
                        original_text,
                        body_style,
                    ),
                    Paragraph(
                        "<b>Modified Text:</b>",
                        body_style,
                    ),
                    Paragraph(
                        modified_text,
                        body_style,
                    ),
                    Paragraph(
                        f"<b>Reason:</b> {reason}",
                        body_style,
                    ),
                ]

                story.append(
                    KeepTogether(
                        modification_block
                    )
                )

        # --------------------------------------------------
        # Validation History
        # --------------------------------------------------

        story.append(
            Paragraph(
                "5. Validation History",
                heading_style,
            )
        )

        history = report_data.get(
            "validation_history",
            [],
        )

        if not history:

            story.append(
                Paragraph(
                    "No validation history available.",
                    body_style,
                )
            )

        else:

            history_data = [
                [
                    Paragraph(
                        "<b>Iteration</b>",
                        small_style,
                    ),
                    Paragraph(
                        "<b>Status</b>",
                        small_style,
                    ),
                    Paragraph(
                        "<b>Risk</b>",
                        small_style,
                    ),
                    Paragraph(
                        "<b>Passed</b>",
                        small_style,
                    ),
                    Paragraph(
                        "<b>Failed</b>",
                        small_style,
                    ),
                ]
            ]

            for iteration in history:

                history_data.append(
                    [
                        Paragraph(
                            self._escape(
                                str(
                                    iteration.get(
                                        "iteration",
                                        "",
                                    )
                                )
                            ),
                            small_style,
                        ),
                        Paragraph(
                            self._escape(
                                iteration.get(
                                    "status",
                                    "",
                                )
                            ),
                            small_style,
                        ),
                        Paragraph(
                            self._escape(
                                iteration.get(
                                    "risk",
                                    "",
                                )
                            ),
                            small_style,
                        ),
                        Paragraph(
                            self._escape(
                                str(
                                    iteration.get(
                                        "passed_rules",
                                        "",
                                    )
                                )
                            ),
                            small_style,
                        ),
                        Paragraph(
                            self._escape(
                                str(
                                    iteration.get(
                                        "failed_rules",
                                        "",
                                    )
                                )
                            ),
                            small_style,
                        ),
                    ]
                )

            history_table = Table(
                history_data,
                colWidths=[
                    30 * mm,
                    35 * mm,
                    30 * mm,
                    35 * mm,
                    35 * mm,
                ],
                repeatRows=1,
            )

            history_table.setStyle(
                TableStyle(
                    [
                        (
                            "GRID",
                            (0, 0),
                            (-1, -1),
                            0.5,
                            colors.grey,
                        ),
                        (
                            "BACKGROUND",
                            (0, 0),
                            (-1, 0),
                            colors.lightgrey,
                        ),
                        (
                            "VALIGN",
                            (0, 0),
                            (-1, -1),
                            "TOP",
                        ),
                        (
                            "LEFTPADDING",
                            (0, 0),
                            (-1, -1),
                            5,
                        ),
                        (
                            "RIGHTPADDING",
                            (0, 0),
                            (-1, -1),
                            5,
                        ),
                        (
                            "TOPPADDING",
                            (0, 0),
                            (-1, -1),
                            4,
                        ),
                        (
                            "BOTTOMPADDING",
                            (0, 0),
                            (-1, -1),
                            4,
                        ),
                    ]
                )
            )

            story.append(
                history_table
            )

        # --------------------------------------------------
        # Build PDF
        # --------------------------------------------------

        document.build(
            story
        )

        return output

    # ======================================================
    # HELPERS
    # ======================================================

    @staticmethod
    def _escape(
        value: Any,
    ) -> str:
        """
        Escape text for ReportLab Paragraph markup.
        """

        text = str(
            value
            if value is not None
            else ""
        )

        return (
            text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )